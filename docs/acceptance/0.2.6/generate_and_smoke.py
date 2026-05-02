#!/usr/bin/env python3

import csv
import json
import shutil
import subprocess
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
FIXTURES = ROOT / "fixtures"
CONVERTED = ROOT / "converted"
MANIFEST = ROOT / "test-manifest.json"
SENTINEL = "SMART_IMPORT_026_SENTINEL"


def reset_dirs():
    for directory in (FIXTURES, CONVERTED):
        if directory.exists():
            shutil.rmtree(directory)
        directory.mkdir(parents=True, exist_ok=True)


def write_text(path, content):
    path.write_text(content, encoding="utf-8")


def make_docx():
    document = Document()
    document.add_heading("Smart Import 0.2.6 DOCX Fixture", level=1)
    document.add_paragraph(f"{SENTINEL} docx body text.")
    document.save(FIXTURES / "smart-import-026-docx.docx")


def make_xlsx():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Acceptance"
    sheet.append(["marker", "format", "value"])
    sheet.append([SENTINEL, "xlsx", 260])
    workbook.save(FIXTURES / "smart-import-026-xlsx.xlsx")


def make_pdf():
    pdf_path = FIXTURES / "smart-import-026-pdf.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 720, "Smart Import 0.2.6 PDF Fixture")
    c.setFont("Helvetica", 12)
    c.drawString(72, 690, f"{SENTINEL} pdf body text.")
    c.save()


def make_pptx():
    pptx_path = FIXTURES / "smart-import-026-pptx.pptx"
    pandoc = shutil.which("pandoc")
    if pandoc:
        source = FIXTURES / "smart-import-026-pptx-source.md"
        write_text(source, f"""% Smart Import 0.2.6 PPTX Fixture

# Smart Import 0.2.6 PPTX Fixture

{SENTINEL} pptx body text.
""")
        result = run_command([pandoc, str(source), "-o", str(pptx_path)])
        source.unlink(missing_ok=True)
        if result.returncode == 0 and pptx_path.exists():
            return

    files = {
        "[Content_Types].xml": """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>
  <Override PartName="/ppt/slides/slide1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>
</Types>""",
        "_rels/.rels": """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/>
</Relationships>""",
        "ppt/presentation.xml": """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:presentation xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <p:sldIdLst><p:sldId id="256" r:id="rId1"/></p:sldIdLst>
  <p:sldSz cx="9144000" cy="6858000" type="screen4x3"/>
</p:presentation>""",
        "ppt/_rels/presentation.xml.rels": """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide1.xml"/>
</Relationships>""",
        "ppt/slides/slide1.xml": f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld>
    <p:spTree>
      <p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr>
      <p:grpSpPr/>
      <p:sp>
        <p:nvSpPr><p:cNvPr id="2" name="Title"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr>
        <p:spPr/>
        <p:txBody><a:bodyPr/><a:lstStyle/><a:p><a:r><a:t>Smart Import 0.2.6 PPTX Fixture</a:t></a:r></a:p></p:txBody>
      </p:sp>
      <p:sp>
        <p:nvSpPr><p:cNvPr id="3" name="Body"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr>
        <p:spPr/>
        <p:txBody><a:bodyPr/><a:lstStyle/><a:p><a:r><a:t>{SENTINEL} pptx body text.</a:t></a:r></a:p></p:txBody>
      </p:sp>
    </p:spTree>
  </p:cSld>
</p:sld>""",
    }
    with zipfile.ZipFile(pptx_path, "w", zipfile.ZIP_DEFLATED) as archive:
      for name, content in files.items():
          archive.writestr(name, content)


def make_epub():
    epub_path = FIXTURES / "smart-import-026-epub.epub"
    with zipfile.ZipFile(epub_path, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        archive.writestr("META-INF/container.xml", """<?xml version="1.0"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles>
</container>""")
        archive.writestr("OEBPS/content.opf", """<?xml version="1.0" encoding="UTF-8"?>
<package version="3.0" unique-identifier="bookid" xmlns="http://www.idpf.org/2007/opf">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:identifier id="bookid">smart-import-026</dc:identifier>
    <dc:title>Smart Import 0.2.6 EPUB Fixture</dc:title>
    <dc:language>zh-CN</dc:language>
  </metadata>
  <manifest><item id="chapter1" href="chapter1.xhtml" media-type="application/xhtml+xml"/></manifest>
  <spine><itemref idref="chapter1"/></spine>
</package>""")
        archive.writestr("OEBPS/chapter1.xhtml", f"""<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml"><head><title>Chapter 1</title></head>
<body><h1>Smart Import 0.2.6 EPUB Fixture</h1><p>{SENTINEL} epub body text.</p></body></html>""")


def make_calibre_ebook_fixtures():
    ebook_convert = shutil.which("ebook-convert")
    epub_path = FIXTURES / "smart-import-026-epub.epub"
    if not ebook_convert or not epub_path.exists():
        return
    for extension in ("mobi", "azw3"):
        output_path = FIXTURES / f"smart-import-026-{extension}.{extension}"
        result = run_command([ebook_convert, str(epub_path), str(output_path)])
        if result.returncode != 0:
            print(f"Failed to generate {extension}: {result.stderr or result.stdout}")


def make_plain_fixtures():
    write_text(FIXTURES / "smart-import-026-md.md", f"# Smart Import 0.2.6 MD Fixture\n\n{SENTINEL} md body text.\n")
    write_text(FIXTURES / "smart-import-026-txt.txt", f"Smart Import 0.2.6 TXT Fixture\n{SENTINEL} txt body text.\n")
    write_text(FIXTURES / "smart-import-026-html.html", f"<!doctype html><html><body><h1>Smart Import 0.2.6 HTML Fixture</h1><p>{SENTINEL} html body text.</p></body></html>\n")
    with (FIXTURES / "smart-import-026-csv.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["marker", "format", "value"])
        writer.writerow([SENTINEL, "csv", 260])
    write_text(FIXTURES / "smart-import-026-json.json", json.dumps({"title": "Smart Import 0.2.6 JSON Fixture", "marker": SENTINEL, "format": "json"}, ensure_ascii=False, indent=2) + "\n")
    write_text(FIXTURES / "smart-import-026-xml.xml", f"""<?xml version="1.0" encoding="UTF-8"?>
<fixture><title>Smart Import 0.2.6 XML Fixture</title><marker>{SENTINEL}</marker></fixture>
""")
    write_text(FIXTURES / "smart-import-026-eml.eml", f"""From: tester@example.com
To: smart-import@example.com
Subject: Smart Import 0.2.6 EML Fixture
Content-Type: text/plain; charset=utf-8

{SENTINEL} eml body text.
""")
    write_text(FIXTURES / "smart-import-026-ipynb.ipynb", json.dumps({
        "cells": [
            {"cell_type": "markdown", "metadata": {}, "source": ["# Smart Import 0.2.6 IPYNB Fixture\n", f"{SENTINEL} ipynb body text."]},
        ],
        "metadata": {},
        "nbformat": 4,
        "nbformat_minor": 5,
    }, ensure_ascii=False, indent=2) + "\n")
    zip_source = FIXTURES / "zip-note.txt"
    write_text(zip_source, f"Smart Import 0.2.6 ZIP Fixture\n{SENTINEL} zip body text.\n")
    with zipfile.ZipFile(FIXTURES / "smart-import-026-zip.zip", "w", zipfile.ZIP_DEFLATED) as archive:
        archive.write(zip_source, arcname="zip-note.txt")
    zip_source.unlink()


def run_command(args):
    return subprocess.run(args, text=True, capture_output=True)


def smoke_convert():
    rows = []
    markitdown = shutil.which("markitdown")
    pandoc = shutil.which("pandoc")
    ebook_convert = shutil.which("ebook-convert")
    for fixture in sorted(FIXTURES.iterdir()):
        ext = fixture.suffix.lower().lstrip(".")
        output = CONVERTED / f"{fixture.stem}.md"
        if ext in {"md", "txt"}:
            content = fixture.read_text(encoding="utf-8", errors="replace")
            output.write_text(content, encoding="utf-8")
            rows.append({"file": fixture.name, "format": ext, "status": "pass", "converter": "direct-copy", "sentinel_found": SENTINEL in content, "detail": ""})
            continue
        if ext in {"mobi", "azw3"} and ebook_convert:
            text_output = CONVERTED / f"{fixture.stem}.txt"
            result = run_command([ebook_convert, str(fixture), str(text_output)])
            converter = "ebook-convert"
            if text_output.exists():
                output.write_text(text_output.read_text(encoding="utf-8", errors="replace"), encoding="utf-8")
        elif ext == "epub" and pandoc:
            result = run_command([pandoc, str(fixture), "-t", "gfm", "-o", str(output)])
            converter = "pandoc"
        elif markitdown:
            result = run_command([markitdown, str(fixture), "-o", str(output)])
            converter = "markitdown"
        else:
            rows.append({"file": fixture.name, "format": ext, "status": "blocked", "converter": "", "sentinel_found": False, "detail": "markitdown not found"})
            continue
        content = output.read_text(encoding="utf-8", errors="replace") if output.exists() else ""
        normalized_content = content.replace("\\_", "_")
        rows.append({
            "file": fixture.name,
            "format": ext,
            "status": "pass" if result.returncode == 0 and SENTINEL in normalized_content else "fail",
            "converter": converter,
            "sentinel_found": SENTINEL in normalized_content,
            "detail": (result.stderr or result.stdout or "").strip()[:500],
        })

    if not ebook_convert:
        rows.extend([
            {"file": "external-mobi-required.mobi", "format": "mobi", "status": "blocked", "converter": "ebook-convert", "sentinel_found": False, "detail": "ebook-convert not found in PATH; install Calibre and rerun with a real MOBI sample."},
            {"file": "external-azw3-required.azw3", "format": "azw3", "status": "blocked", "converter": "ebook-convert", "sentinel_found": False, "detail": "ebook-convert not found in PATH; install Calibre and rerun with a real AZW3 sample."},
        ])
    rows.append({"file": "external-msg-required.msg", "format": "msg", "status": "manual", "converter": "markitdown", "sentinel_found": False, "detail": "Use a real Outlook .msg sample for UI acceptance; plain-text fixtures cannot represent MSG binary structure."})
    return rows


def write_manifest(rows):
    payload = {
        "plugin_version": "0.2.6",
        "sentinel": SENTINEL,
        "dependencies": {
            "markitdown": shutil.which("markitdown") or "",
            "pandoc": shutil.which("pandoc") or "",
            "ebook-convert": shutil.which("ebook-convert") or "",
        },
        "results": rows,
    }
    MANIFEST.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main():
    reset_dirs()
    make_plain_fixtures()
    make_docx()
    make_xlsx()
    make_pdf()
    make_pptx()
    make_epub()
    make_calibre_ebook_fixtures()
    rows = smoke_convert()
    write_manifest(rows)
    passed = sum(1 for row in rows if row["status"] == "pass")
    failed = sum(1 for row in rows if row["status"] == "fail")
    blocked = sum(1 for row in rows if row["status"] == "blocked")
    manual = sum(1 for row in rows if row["status"] == "manual")
    print(json.dumps({"passed": passed, "failed": failed, "blocked": blocked, "manual": manual, "manifest": str(MANIFEST)}, ensure_ascii=False))
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
